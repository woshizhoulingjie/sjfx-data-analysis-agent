import json
import logging
import os
import re
import shutil
import zipfile
import hashlib
import time
import uuid
from contextlib import contextmanager, suppress
from datetime import datetime
from pathlib import Path, PurePosixPath

from services.evidence import select_evidence
from services.unified_parser import SUPPORTED_EXTENSIONS


LOGGER = logging.getLogger(__name__)

# Handoff exports intentionally keep the historical secret/noise exclusions.
# Scanner inventory uses a different, explicit policy and must never inherit
# these defaults, otherwise excluded entries disappear from coverage.
EXPORT_IGNORED_DIRS = {
    ".git", ".venv", ".venv_py37_unused", "__pycache__", "node_modules",
    ".idea", ".vscode", "vendor_packages", "work",
}
EXPORT_IGNORED_FILES = {".env", "agent.db", "agent.db-shm", "agent.db-wal"}
EXPORT_SENSITIVE_EXTENSIONS = {".key", ".pem", ".p12", ".pfx", ".keystore"}


def _should_ignore_export_file(name):
    lower_name = str(name or "").casefold()
    return (
        lower_name in EXPORT_IGNORED_FILES
        or lower_name.startswith("~$")
        or Path(lower_name).suffix in EXPORT_SENSITIVE_EXTENSIONS
    )


def xml_safe_text(value):
    """Return text that is safe for XML 1.0 / python-docx.

    Parsed files can legally contain NUL bytes, terminal control characters,
    Unicode noncharacters, or even isolated UTF-16 surrogate code points.  XML
    1.0 cannot represent those values and lxml raises before the report can be
    saved.  Keep normal text plus tab/newline/carriage return, and filter only
    characters that are unsafe for the document package.
    """
    if value is None:
        return ""
    text = str(value)
    safe = []
    for character in text:
        codepoint = ord(character)
        if codepoint in (0x09, 0x0A, 0x0D):
            safe.append(character)
            continue
        if codepoint < 0x20 or 0x7F <= codepoint <= 0x9F:
            continue
        if 0xD800 <= codepoint <= 0xDFFF:
            continue
        if codepoint > 0x10FFFF:
            continue
        # Unicode noncharacters cannot carry meaningful source text and are
        # rejected by several XML consumers even where a parser is permissive.
        if 0xFDD0 <= codepoint <= 0xFDEF:
            continue
        if (codepoint & 0xFFFF) in (0xFFFE, 0xFFFF):
            continue
        safe.append(character)
    return "".join(safe)


def safe_name(value):
    return re.sub(r"[^\w\-.\u4e00-\u9fff]+", "_", value, flags=re.UNICODE).strip("_") or "export"


def collect_files(path):
    path = Path(path)
    if path.is_symlink():
        return []
    if path.is_file():
        return [path]
    files = []
    for current_root, dirs, names in os.walk(str(path)):
        dirs[:] = [d for d in dirs if d.casefold() not in EXPORT_IGNORED_DIRS and not (Path(current_root) / d).is_symlink()]
        for name in names:
            if _should_ignore_export_file(name):
                continue
            candidate = Path(current_root) / name
            if candidate.is_symlink():
                continue
            files.append(candidate)
    return files


def _collect_inventoried_files(root, selected, inventory_metadata):
    """Return lexical source paths from the immutable scan inventory.

    A live ``os.walk`` silently skips a regular file that has been replaced by
    a symlink. Building the export set from the recorded inventory ensures the
    handle-level opener sees that replacement and rejects the export.
    """
    root = Path(root).resolve()
    selected = Path(selected)
    try:
        selected_relative = selected.relative_to(root)
    except ValueError as exc:
        raise ValueError("导出范围超出扫描根目录") from exc
    selected_key = str(selected_relative).replace("\\", "/") or "."
    selected_prefix = "" if selected_key == "." else selected_key.rstrip("/") + "/"
    files = []
    excluded = []
    for metadata_key, raw_node in sorted(
        (inventory_metadata or {}).items(), key=lambda item: str(item[0])
    ):
        node = raw_node if isinstance(raw_node, dict) else {}
        if node.get("kind") != "file":
            continue
        relative = str(node.get("path") or metadata_key or "").replace("\\", "/").strip("/")
        parts = PurePosixPath(relative).parts
        if not relative or relative == "." or ".." in parts or PurePosixPath(relative).is_absolute():
            raise ValueError("清点记录包含非法导出路径")
        if selected_key != "." and relative != selected_key and not relative.startswith(selected_prefix):
            continue
        candidate = root.joinpath(*parts)
        ignored = (
            any(part.casefold() in EXPORT_IGNORED_DIRS for part in parts[:-1])
            or _should_ignore_export_file(parts[-1])
        )
        if ignored:
            excluded.append(candidate)
        else:
            files.append(candidate)
    return files, excluded


def _sha256_file(path, block_size=1024 * 1024):
    """Return the raw-source digest used for handoff-package deduplication."""
    digest = hashlib.sha256()
    with Path(path).open("rb") as handle:
        while True:
            block = handle.read(block_size)
            if not block:
                break
            digest.update(block)
    return digest.hexdigest()


def _inventory_node_for_source(root, path, inventory_metadata=None):
    root = Path(root).resolve()
    path = Path(path)
    try:
        relative_path = path.relative_to(root)
    except ValueError as exc:
        raise ValueError("导出源文件超出扫描根目录") from exc
    if ".." in relative_path.parts:
        raise ValueError("导出源文件路径包含越界分量")
    relative = str(relative_path).replace("\\", "/")
    node = dict((inventory_metadata or {}).get(relative) or {})
    if node:
        return relative, node
    # Compatibility for direct library callers: create a one-shot inventory
    # record with lstat, then the handle-level opener below verifies it again.
    stat = path.lstat()
    if path.is_symlink():
        raise ValueError("导出源文件不能是符号链接")
    return relative, {
        "path": relative,
        "name": path.name,
        "extension": path.suffix.lower(),
        "size": int(stat.st_size),
        "modified_at_ns": int(stat.st_mtime_ns),
        "device": int(stat.st_dev),
        "inode": int(stat.st_ino),
    }


@contextmanager
def _open_verified_export_source(root, path, inventory_metadata=None):
    from services.package_analysis import _open_inventory_source

    relative, node = _inventory_node_for_source(root, path, inventory_metadata)
    descriptor, opened = _open_inventory_source(root, node)
    try:
        stream = os.fdopen(descriptor, "rb", closefd=True)
    except Exception:
        os.close(descriptor)
        raise
    with stream:
        yield relative, node, opened, stream


def _deduplicate_files(files, root, cancel_check=None, inventory_metadata=None):
    """Keep one stable canonical source per identical byte stream.

    UI selections are path based, but a handoff package must not contain two
    distinct names for exactly the same source bytes.  The manifest retains all
    omitted paths, so no provenance is lost.
    """
    groups = {}
    for path in sorted(set(files), key=lambda item: str(item.relative_to(root)).replace("\\", "/")):
        if cancel_check:
            cancel_check()
        digest_builder = hashlib.sha256()
        with _open_verified_export_source(root, path, inventory_metadata) as (_relative, _node, before, stream):
            while True:
                if cancel_check:
                    cancel_check()
                block = stream.read(4 * 1024 * 1024)
                if not block:
                    break
                digest_builder.update(block)
            after = os.fstat(stream.fileno())
        if (
            int(before.st_dev), int(before.st_ino), int(before.st_size), int(before.st_mtime_ns)
        ) != (
            int(after.st_dev), int(after.st_ino), int(after.st_size), int(after.st_mtime_ns)
        ):
            raise ValueError("源文件在导出去重期间发生变化：{}".format(path.name))
        digest = digest_builder.hexdigest()
        groups.setdefault(digest, []).append(path)

    unique_files = []
    duplicates = []
    for digest, members in groups.items():
        canonical = members[0]
        unique_files.append(canonical)
        if len(members) > 1:
            duplicates.append({
                "sha256": digest,
                "canonical": str(canonical.relative_to(root)).replace("\\", "/"),
                "omitted": [str(item.relative_to(root)).replace("\\", "/") for item in members[1:]],
            })
    return unique_files, duplicates, {
        member: digest for digest, members in groups.items() for member in members
    }


@contextmanager
def _atomic_zip(temporary_path, final_path):
    """Write a ZIP privately and publish it only after it closes cleanly."""
    temporary_path = Path(temporary_path)
    final_path = Path(final_path)
    temporary_path.parent.mkdir(parents=True, exist_ok=True)
    try:
        with zipfile.ZipFile(str(temporary_path), "w", compression=zipfile.ZIP_DEFLATED,
                             allowZip64=True) as archive:
            yield archive
        if os.name != "nt":
            temporary_path.chmod(0o600)
        # Ensure the directory entry is durable before publishing the name.
        # A crash can therefore leave only a harmless .part file, never a
        # half-valid final archive.
        try:
            with temporary_path.open("rb") as stream:
                os.fsync(stream.fileno())
        except (OSError, AttributeError):
            pass
        os.replace(str(temporary_path), str(final_path))
    except BaseException:
        with suppress(OSError):
            temporary_path.unlink()
        raise


def cleanup_stale_part_files(output_dir, max_age_seconds=24 * 60 * 60):
    """Remove abandoned exporter work files, never published archives.

    Only files directly under the configured output directory with the exact
    ``*.zip.part`` suffix are considered, and recent files are preserved in
    case another process is still writing them.  This makes recovery after a
    worker kill deterministic without risking a valid ``.zip`` result.
    """
    output_dir = Path(output_dir)
    if not output_dir.exists() or not output_dir.is_dir() or output_dir.is_symlink():
        return 0
    now = time.time()
    removed = 0
    try:
        candidates = output_dir.glob("*.zip.part")
        for item in candidates:
            try:
                if not item.is_file() or item.is_symlink():
                    continue
                age = max(0.0, now - item.stat().st_mtime)
                if age < max(60, int(max_age_seconds)):
                    continue
                item.unlink()
                removed += 1
            except OSError:
                LOGGER.warning("清理导出临时文件失败：%s", item, exc_info=True)
    except OSError:
        LOGGER.warning("扫描导出临时文件失败：%s", output_dir, exc_info=True)
    return removed


def export_node(root, selected, summary, output_dir, max_bytes, analysis=None, documents=None, task_topic=None,
                member_paths=None, node_name=None, node_id=None, selection_metadata=None,
                selected_evidence_ids=None, inventory_metadata=None, file_states=None,
                progress_callback=None, cancel_check=None, content_deduplication=True,
                known_hashes=None, disk_reserve_bytes=1024 * 1024 * 1024,
                _allow_segmented=True):
    root = Path(root).resolve()
    selected = Path(selected).resolve()
    documents = documents or []
    analysis = analysis or {}
    selection_metadata = selection_metadata or []
    selected_evidence_ids = set(selected_evidence_ids or [])
    inventory_metadata = inventory_metadata or {}
    file_states = file_states or {}
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    cleanup_stale_part_files(output_dir)
    virtual_paths = sorted(set(member_paths or []))
    if virtual_paths:
        selected_rel = "virtual:{}".format(node_id or safe_name(node_name or "topic"))
        files = []
        excluded_files = []
        for relative_path in virtual_paths:
            candidate = root / relative_path
            try:
                candidate.relative_to(root)
            except ValueError:
                raise ValueError("主题节点包含越界文件，已拒绝导出")
            if ".." in candidate.relative_to(root).parts or (
                inventory_metadata and str(relative_path).replace("\\", "/") not in inventory_metadata
            ):
                excluded_files.append(candidate)
            else:
                # Handoff packages preserve every selected source file. Parsing
                # support controls analysis coverage, never raw-material export.
                files.append(candidate)
        selected_documents = [item for item in documents if item.get("path") in set(virtual_paths)]
        export_label = node_name or "主题节点"
    else:
        if inventory_metadata:
            all_files, excluded_files = _collect_inventoried_files(
                root, selected, inventory_metadata
            )
        else:
            all_files = collect_files(selected)
            excluded_files = []
        # Export all original files, including formats the parser cannot read.
        # Unsupported files remain visible in the coverage manifest as metadata-only.
        files = all_files
        selected_rel = str(selected.relative_to(root)).replace("\\", "/") if selected != root else "."
        if selected_rel == ".":
            selected_documents = documents
        else:
            prefix = selected_rel.rstrip("/") + "/"
            selected_documents = [item for item in documents if item.get("path") == selected_rel or item.get("path", "").startswith(prefix)]
        export_label = selected.name
    selected_file_count = len(files)
    # Cheap metadata and capacity gates must run before any SHA-256 pass.  This
    # prevents a doomed 10 GiB export from reading the NAS twice merely to
    # report that it is too large or the destination is full.
    source_stats = {}
    try:
        for path in files:
            with _open_verified_export_source(root, path, inventory_metadata) as (relative, node, stat, _stream):
                source_stats[path] = {
                    "relative": relative,
                    "node": node,
                    "size": int(stat.st_size),
                    "mtime_ns": int(stat.st_mtime_ns),
                }
    except (OSError, PermissionError) as exc:
        raise ValueError("导出前无法读取源文件：{}".format(exc))
    total_size = sum(item["size"] for item in source_stats.values())
    if total_size > max_bytes and not _allow_segmented:
        raise ValueError("单个分卷内容超过 {:.2f} GB 上限".format(max_bytes / 1073741824))
    try:
        free_bytes = int(shutil.disk_usage(str(output_dir)).free)
    except OSError as exc:
        raise ValueError("无法检查导出磁盘剩余空间：{}".format(exc))
    required_bytes = total_size + max(0, int(disk_reserve_bytes or 0))
    if free_bytes < required_bytes:
        raise ValueError(
            "导出磁盘空间不足：最坏情况需要 {:.2f} GB（含保留空间），当前可用 {:.2f} GB".format(
                required_bytes / 1073741824, free_bytes / 1073741824,
            )
        )
    if total_size > max_bytes:
        volumes = []
        current = []
        current_size = 0
        for path in sorted(files, key=lambda item: str(item.relative_to(root)).replace("\\", "/")):
            size = source_stats[path]["size"]
            if size > max_bytes:
                raise ValueError("单文件 {} 为 {:.2f} GB，超过单卷上限，无法安全分卷".format(path.name, size / 1073741824))
            if current and current_size + size > max_bytes:
                volumes.append(current)
                current = []
                current_size = 0
            current.append(path)
            current_size += size
        if current:
            volumes.append(current)

        generated_parts = []
        completed_files = 0
        for index, volume in enumerate(volumes, 1):
            if cancel_check:
                cancel_check()
            relative_paths = [str(path.relative_to(root)).replace("\\", "/") for path in volume]
            part_path = export_node(
                root, selected, summary, output_dir, max_bytes,
                analysis=analysis, documents=documents, task_topic=task_topic,
                member_paths=relative_paths,
                node_name="{}_分卷{:03d}".format(export_label, index),
                node_id="{}-part-{:03d}".format(node_id or safe_name(export_label), index),
                selection_metadata=selection_metadata,
                selected_evidence_ids=selected_evidence_ids,
                inventory_metadata=inventory_metadata, file_states=file_states,
                progress_callback=None, cancel_check=cancel_check,
                content_deduplication=content_deduplication,
                known_hashes=known_hashes, disk_reserve_bytes=disk_reserve_bytes,
                _allow_segmented=False,
            )
            generated_parts.append({
                "index": index, "file_name": part_path.name,
                "size": part_path.stat().st_size, "sha256": _sha256_file(part_path),
                "source_file_count": len(volume), "source_paths": relative_paths,
            })
            completed_files += len(volume)
            if progress_callback:
                progress_callback(completed_files, len(files), sum(item["size"] for item in generated_parts), total_size)

        stamp = datetime.now().strftime("%Y%m%d-%H%M%S")
        index_path = Path(output_dir) / "待整编数据包_{}_{}_{}_分卷索引.zip".format(
            safe_name(export_label), stamp, uuid.uuid4().hex[:8]
        )
        manifest = {
            "schema_version": "segmented-export/1.0",
            "task_topic": str(task_topic or ""),
            "source_total_bytes": total_size,
            "source_file_count": len(files),
            "volume_limit_bytes": max_bytes,
            "volume_count": len(generated_parts),
            "parts": generated_parts,
            "instructions": "逐个下载所有分卷；每个分卷都是可独立打开的 ZIP，分卷清单中的 SHA-256 用于完整性核验。",
        }
        with _atomic_zip(index_path.with_name(index_path.name + ".part"), index_path) as archive:
            archive.writestr("分卷清单.json", json.dumps(manifest, ensure_ascii=False, indent=2))
            archive.writestr("README.txt", manifest["instructions"])
        sidecar = index_path.with_name(index_path.name + ".parts.json")
        sidecar.write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
        if os.name != "nt":
            sidecar.chmod(0o600)
        return index_path
    expected_hashes = {}
    content_duplicates = []
    if content_deduplication:
        try:
            files, content_duplicates, expected_hashes = _deduplicate_files(
                files, root, cancel_check=cancel_check,
                inventory_metadata=inventory_metadata,
            )
        except (OSError, PermissionError) as exc:
            raise ValueError("导出前无法计算源文件去重指纹：{}".format(exc))
    known_hashes = {str(key).replace("\\", "/"): str(value) for key, value in (known_hashes or {}).items() if value}
    for path in files:
        relative = str(path.relative_to(root)).replace("\\", "/")
        expected_hashes.setdefault(path, known_hashes.get(relative))
    total_size = sum(source_stats[path]["size"] for path in files)
    stamp = datetime.now().strftime("%Y%m%d-%H%M%S")
    # A random suffix prevents two same-second exports from sharing the same
    # final or .part pathname.  The human-readable prefix remains unchanged.
    archive_name = "待整编数据包_{}_{}_{}.zip".format(
        safe_name(export_label), stamp, uuid.uuid4().hex[:8]
    )
    archive_path = Path(output_dir) / archive_name
    temporary_archive_path = archive_path.with_name(archive_path.name + ".part")
    evidence_candidates = []
    for item in selected_documents:
        evidence_candidates.extend(item.get("payload", {}).get("evidence", []))
    if selected_evidence_ids:
        evidence_candidates = [
            item for item in evidence_candidates
            if item.get("evidence_id") in selected_evidence_ids
        ]
    summary_topics = list((summary or {}).get("topics", []))
    direction = (summary or {}).get("recommended_research_direction", {})
    summary_topics.extend([direction.get("title"), direction.get("rationale")])
    evidence_chain = select_evidence(
        evidence_candidates, topics=summary_topics, max_items=40, per_source=2, max_chars=600
    )
    task_topic = str(task_topic or "").strip() or None
    if not task_topic:
        raise ValueError("导出前必须指定整编任务主题")
    handoff = {
        "schema_version": "compilation-task/2.0",
        "selected_path": selected_rel,
        "selected_node_name": node_name,
        "selected_node_id": node_id,
        "task_topic": task_topic,
        "task_topic_required": True,
        "task_topic_status": "provided" if task_topic else "customer_input_required",
        "instruction": (
            "整编 Agent 必须围绕 task_topic 完成写作；若其为 null，必须先要求客户指定主题，"
            "不得把节点摘要或推荐研究方向自动当作客户任务主题。"
        ),
        "selected_node_summary": (summary or {}).get("summary"),
        "recommended_research_direction": (summary or {}).get("recommended_research_direction", {}),
        "analysis_coverage": analysis.get("coverage", {}),
        "value_judgment": analysis.get("value_judgment", {}),
        "evidence_file": "结论-证据链.json",
        "coverage_file": "解析覆盖率清单.json",
        "selection_mode": "combined" if len(selection_metadata) > 1 else "single",
        "selected_nodes": selection_metadata,
        "unique_source_file_count": len(files),
        "deduplication": {
            "method": (
                "先按相对路径合并重叠选择，再按 SHA-256 精确去重；写入时再次校验摘要。"
                if content_deduplication else
                "10GB/大包流式导出跳过内容去重；按路径去重，写入 ZIP 时单遍计算 SHA-256。"
            ),
            "content_deduplication_enabled": bool(content_deduplication),
            "source_selection_count": len(selection_metadata),
            "selected_evidence_count": len(selected_evidence_ids),
            "selected_file_count_before_content_deduplication": selected_file_count,
            "content_duplicate_group_count": len(content_duplicates),
            "content_duplicate_file_count": sum(len(item["omitted"]) for item in content_duplicates),
            "content_duplicate_groups": content_duplicates,
        },
    }
    written_size = 0
    source_manifest = []
    with _atomic_zip(temporary_archive_path, archive_path) as archive:
        for index, path in enumerate(files, 1):
            if cancel_check:
                cancel_check()
            expected_stat = source_stats[path]
            expected_size = expected_stat["size"]
            if written_size + expected_size > max_bytes:
                raise ValueError("源文件总大小超过导出上限，已停止写入")
            digest = hashlib.sha256()
            copied = 0
            with _open_verified_export_source(root, path, inventory_metadata) as (relative, _node, before, source):
                if int(before.st_size) != expected_size or int(before.st_mtime_ns) != expected_stat["mtime_ns"]:
                    raise ValueError("源文件在导出过程中发生变化：{}".format(path.name))
                info = zipfile.ZipInfo(relative, time.localtime(before.st_mtime)[:6])
                info.compress_type = zipfile.ZIP_DEFLATED
                info.external_attr = (int(before.st_mode) & 0xFFFF) << 16
                with archive.open(info, "w", force_zip64=True) as target:
                    while True:
                        if cancel_check:
                            cancel_check()
                        block = source.read(4 * 1024 * 1024)
                        if not block:
                            break
                        copied += len(block)
                        if written_size + copied > max_bytes:
                            raise ValueError("源文件总大小超过导出上限，已停止写入")
                        digest.update(block)
                        target.write(block)
                after = os.fstat(source.fileno())
            actual_sha256 = digest.hexdigest()
            if (
                copied != expected_size
                or int(after.st_size) != expected_size
                or int(after.st_mtime_ns) != expected_stat["mtime_ns"]
            ):
                raise ValueError("源文件在导出过程中发生变化：{}".format(path.name))
            expected_sha256 = expected_hashes.get(path)
            if expected_sha256 and actual_sha256 != expected_sha256:
                raise ValueError("源文件内容与分析/去重摘要不一致：{}".format(path.name))
            written_size += copied
            source_manifest.append({
                "path": relative, "size": copied, "sha256": actual_sha256,
                "modified_at_ns": expected_stat["mtime_ns"],
            })
            if progress_callback:
                progress_callback(index, len(files), written_size, total_size)
        handoff["source_manifest_file"] = "源文件SHA-256清单.json"
        handoff["source_manifest_count"] = len(source_manifest)
        archive.writestr(
            "节点摘要.json",
            json.dumps(summary or {"message": "尚未生成摘要"}, ensure_ascii=False, indent=2),
        )
        archive.writestr("整编任务说明.json", json.dumps(handoff, ensure_ascii=False, indent=2))
        archive.writestr("源文件SHA-256清单.json", json.dumps({
            "schema_version": "source-manifest/1.0",
            "total_bytes": written_size,
            "files": source_manifest,
        }, ensure_ascii=False, indent=2))
        archive.writestr("结论-证据链.json", json.dumps({
            "schema_version": "question-answer-evidence/2.0",
            "selected_path": selected_rel,
            "selected_node_name": node_name,
            "conclusions": (summary or {}).get("conclusion_evidence", []),
            "question_answer_evidence": (summary or {}).get("question_answer_evidence") or (summary or {}).get("conclusion_evidence", []),
            "claims": (summary or {}).get("claims", []),
            "evidence_status": (summary or {}).get("evidence_status") or ("supported" if evidence_chain else "insufficient"),
            "evidence_count": len(evidence_chain),
            "candidate_count": len(evidence_candidates),
            "omitted_count": max(0, len(evidence_candidates) - len(evidence_chain)),
            "selection_method": "主题相关度 + 页/节可追溯性 + 每文件最多 2 条",
            "topic_terms": summary_topics[:12],
            "items": evidence_chain,
        }, ensure_ascii=False, indent=2))
        archive.writestr("统一文档索引.json", json.dumps([
            {
                "path": item.get("path"),
                "schema_version": item.get("payload", {}).get("schema_version"),
                "source": item.get("payload", {}).get("source"),
                "parser": item.get("payload", {}).get("parser"),
                "structure": item.get("payload", {}).get("structure"),
                "content_sha256": item.get("payload", {}).get("content_sha256"),
                "warnings": item.get("payload", {}).get("warnings", []),
            } for item in selected_documents
        ], ensure_ascii=False, indent=2))
        archive.writestr("去重与聚类清单.json", json.dumps({
            "statistics": analysis.get("statistics", {}),
            "exact_duplicate_groups": analysis.get("exact_duplicate_groups", []),
            "similar_document_clusters": analysis.get("similar_document_clusters", []),
            "topic_clusters": analysis.get("topic_clusters", []),
            "classification_dimensions": analysis.get("classification_dimensions", []),
        }, ensure_ascii=False, indent=2))
        archive.writestr("检索证据.json", json.dumps({
            "schema_version": "local-retrieval/1.0",
            "selected_path": selected_rel,
            "retrieval": analysis.get("retrieval", {}),
            "research_retrieval": analysis.get("research_retrieval", {}),
        }, ensure_ascii=False, indent=2))
        document_by_path = {item.get("path"): item.get("payload", {}) for item in selected_documents}
        coverage_items = []
        for relative_path in sorted(set(virtual_paths) if virtual_paths else {
            str(path.relative_to(root)).replace("\\", "/") for path in files
        }):
            payload = document_by_path.get(relative_path, {})
            meta = inventory_metadata.get(relative_path, {})
            state = file_states.get(relative_path, {})
            coverage_items.append({
                "path": relative_path,
                "source": payload.get("source") or meta,
                "analysis_status": state.get("status") or ("completed" if payload else "not_parsed"),
                "parser": payload.get("parser"),
                "structure": payload.get("structure"),
                "stored_characters": len(payload.get("text", "")),
                "evidence_count": len(payload.get("evidence", [])),
                "coverage": payload.get("coverage", {}),
                "warnings": payload.get("warnings", []),
            })
        archive.writestr("解析覆盖率清单.json", json.dumps({
            "schema_version": "parse-coverage/1.0",
            "selected_path": selected_rel,
            "file_count": len(coverage_items),
            "parsed_file_count": sum(1 for item in coverage_items if item["analysis_status"] in {"completed", "overview"}),
            "coverage": (analysis.get("coverage") or {}),
            "value_judgment": (analysis.get("value_judgment") or {}),
            "analysis_limitations": list((analysis.get("coverage") or {}).get("limitations") or []),
            "items": coverage_items,
        }, ensure_ascii=False, indent=2))
        archive.writestr(
            "导出清单.json",
            json.dumps({
                "selected_path": selected_rel,
                "file_count": len(files),
                "selected_file_count_before_content_deduplication": selected_file_count,
                "unique_source_file_count": len(files),
                "content_duplicate_group_count": len(content_duplicates),
                "content_duplicate_file_count": sum(len(item["omitted"]) for item in content_duplicates),
                "content_duplicate_groups": content_duplicates,
                "selected_node_count": len(selection_metadata),
                "excluded_file_count": len(excluded_files),
                "excluded_files": [str(path.relative_to(root)).replace("\\", "/") for path in excluded_files[:100]],
                "total_size": total_size,
                "generated_at": datetime.now().isoformat(timespec="seconds"),
                "contents": ["所选范围全部原始文件", "源文件SHA-256清单.json", "节点摘要.json", "整编任务说明.json", "结论-证据链.json", "统一文档索引.json", "去重与聚类清单.json", "检索证据.json", "解析覆盖率清单.json"],
            }, ensure_ascii=False, indent=2),
        )
    return archive_path


def create_report_docx(report, scan, output_path):
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Inches, Pt, RGBColor
    except ImportError:
        raise RuntimeError("未安装 python-docx，无法生成 Word 报告")

    def add_paragraph(text="", style=None):
        return doc.add_paragraph(xml_safe_text(text), style=style)

    def add_heading(text, level):
        return doc.add_heading(xml_safe_text(text), level=level)

    def add_run(paragraph, text=""):
        return paragraph.add_run(xml_safe_text(text))

    def set_xml_text(element, value):
        element.text = xml_safe_text(value)

    def set_font(run, name="Microsoft YaHei", size=None, color=None, bold=None):
        run.font.name = name
        run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
        run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
        run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
        if size is not None:
            run.font.size = Pt(size)
        if color is not None:
            run.font.color.rgb = RGBColor(*color)
        if bold is not None:
            run.bold = bold

    def clean_evidence(value, limit=220):
        text = re.sub(r"\s+", " ", xml_safe_text(value)).replace("|", " / ").strip(" /-")
        return text[:limit] + ("…" if len(text) > limit else "")

    def configure_style(style, size, color, before, after, line_spacing, bold=False):
        style.font.name = "Microsoft YaHei"
        style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor(*color)
        style.font.bold = bold
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = line_spacing

    def add_items(values, style="List Bullet", empty_text="暂无可用信息"):
        values = values or []
        if not values:
            add_paragraph(empty_text)
            return
        for value in values:
            add_paragraph(value, style=style)

    def add_conclusion_evidence(values):
        for conclusion in values or []:
            if not isinstance(conclusion, dict):
                continue
            statement = conclusion.get("statement") or "分析结论"
            confidence = conclusion.get("confidence") or "待核验"
            add_paragraph("结论：{}（置信度：{}）".format(statement, confidence), style="List Bullet")
            if conclusion.get("basis"):
                add_paragraph("依据：{}".format(conclusion["basis"]))
            for evidence in conclusion.get("evidence", [])[:4]:
                if not isinstance(evidence, dict):
                    continue
                location = evidence.get("source_path", "未知文件")
                if evidence.get("page"):
                    location += "，第 {} 页".format(evidence["page"])
                add_paragraph(
                    "[{}] {}：{}".format(
                        evidence.get("evidence_id") or "证据",
                        location,
                        clean_evidence(evidence.get("text")),
                    ),
                    style="List Bullet",
                )

    def add_table(headers, rows, widths=None):
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        table.autofit = True
        for index, value in enumerate(headers):
            cell = table.rows[0].cells[index]
            cell.text = xml_safe_text(value)
            for run in cell.paragraphs[0].runs:
                set_font(run, size=9, bold=True, color=(31, 77, 120))
        for values in rows:
            cells = table.add_row().cells
            for index, value in enumerate(values):
                cells[index].text = xml_safe_text(value)
                for paragraph in cells[index].paragraphs:
                    for run in paragraph.runs:
                        set_font(run, size=8.5)
        doc.add_paragraph().paragraph_format.space_after = Pt(1)
        return table

    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    configure_style(styles["Normal"], 11, (0, 0, 0), 0, 6, 1.10)
    configure_style(styles["Heading 1"], 16, (46, 116, 181), 16, 8, 1.10, True)
    configure_style(styles["Heading 2"], 13, (46, 116, 181), 12, 6, 1.10, True)
    configure_style(styles["Heading 3"], 12, (31, 77, 120), 8, 4, 1.10, True)
    configure_style(styles["List Bullet"], 11, (0, 0, 0), 0, 8, 1.167)
    configure_style(styles["List Number"], 11, (0, 0, 0), 0, 8, 1.167)

    header = section.header.paragraphs[0]
    set_xml_text(header, "数据分析 Agent｜数据包情况概览")
    set_font(header.runs[0], size=9, color=(100, 100, 100))
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = add_run(footer, "第 ")
    set_font(footer_run, size=9, color=(100, 100, 100))
    field_begin = OxmlElement("w:fldChar")
    field_begin.set(qn("w:fldCharType"), "begin")
    field_instr = OxmlElement("w:instrText")
    field_instr.set(qn("xml:space"), "preserve")
    set_xml_text(field_instr, " PAGE ")
    field_end = OxmlElement("w:fldChar")
    field_end.set(qn("w:fldCharType"), "end")
    footer_run._r.extend([field_begin, field_instr, field_end])
    tail = add_run(footer, " 页")
    set_font(tail, size=9, color=(100, 100, 100))

    title = add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    title_run = add_run(title, "数据包情况概览报告")
    set_font(title_run, size=23, bold=True)
    subtitle = add_paragraph("自动扫描、证据摘要与研究方向建议")
    subtitle.paragraph_format.space_after = Pt(16)
    set_font(subtitle.runs[0], size=13, color=(70, 70, 70))

    for label, value in (
        ("扫描目录", scan["root"]),
        ("统计范围", "递归文件 {} 个；子目录 {} 个；总大小 {}".format(scan["file_count"], scan.get("directory_count", 0), scan["total_size_human"])),
        ("扫描时间", scan["scanned_at"]),
        ("生成模式", "模型分析（解析与证据链仍为本地）" if report.get("generation_mode") == "model_analyzed" else "本地解析完成，研究方向待模型分析"),
    ):
        paragraph = add_paragraph()
        paragraph.paragraph_format.space_after = Pt(2)
        label_run = add_run(paragraph, label + "：")
        set_font(label_run, size=10.5, bold=True)
        value_run = add_run(paragraph, value)
        set_font(value_run, size=10.5)

    doc.add_page_break()
    add_heading("一、数据包基本信息", level=1)
    add_items(report.get("basic_information"))
    coverage = report.get("coverage") or {}
    if coverage:
        add_heading("覆盖等级与限制", level=2)
        add_table(
            ["指标", "已完成", "总量", "占比/状态"],
            [
                ["目录清点", coverage.get("scanned_files", 0), coverage.get("inventory_files", 0), "完整" if (coverage.get("inventory_coverage") or {}).get("complete") else "存在限制"],
                ["内容解析", coverage.get("parsed_files", 0), coverage.get("inventory_files", 0), "{:.1%}".format(float(coverage.get("content_parse_ratio") or 0))],
                ["全文深度分析", coverage.get("deep_analyzed_files", 0), coverage.get("inventory_files", 0), "{:.1%}".format(float(coverage.get("deep_analysis_ratio") or 0))],
                ["失败/待处理", coverage.get("failed_files", 0), coverage.get("pending_files", 0), coverage.get("status", "待分析")],
            ],
        )
        add_paragraph(
            "{}（{}）：已解析 {}/{} 个文件；抽样 {} 个；待处理 {} 个；失败 {} 个。".format(
                coverage.get("coverage_level_label", "覆盖等级未标注"),
                coverage.get("status", "待分析"),
                coverage.get("parsed_files", 0),
                coverage.get("inventory_files", 0),
                coverage.get("sampled_files", 0),
                coverage.get("pending_files", 0),
                coverage.get("failed_files", 0),
            )
        )
        archive_totals = coverage.get("archive_member_totals") or {}
        if archive_totals.get("total_members"):
            add_paragraph(
                "压缩包成员：已解析 {}/{}；跳过 {}；失败 {}。".format(
                    archive_totals.get("parsed_members", 0), archive_totals.get("total_members", 0),
                    archive_totals.get("skipped_members", 0), archive_totals.get("failed_members", 0),
                )
            )
        add_items(coverage.get("limitations"), empty_text="当前未发现覆盖限制。")
    judgment = report.get("value_judgment") or {}
    if judgment:
        add_heading("数据价值判断", level=2)
        labels = (
            ("data_usability", "数据可用性"),
            ("information_richness", "信息丰富度"),
            ("research_potential", "研究潜力"),
            ("task_relevance", "与客户任务的相关性"),
        )
        for key, label in labels:
            item = judgment.get(key) or {}
            add_paragraph(
                "{}：{}{}。依据：{}".format(
                    label,
                    item.get("level", "未评估"),
                    "（{}分）".format(item["score"]) if item.get("score") is not None else "",
                    item.get("basis", "未提供"),
                ),
                style="List Bullet",
            )
        add_paragraph(
            "规范文档 {} 份；重复副本 {} 份；有效正文证据 {} 条。".format(
                judgment.get("canonical_document_count", 0), judgment.get("duplicate_alias_count", 0),
                judgment.get("valid_evidence_count", 0),
            )
        )
        add_items(judgment.get("limitations"), empty_text="当前没有额外价值判断限制。")
    intelligence = report.get("intelligence_overview") or {}
    if intelligence:
        add_heading("情报概览指标", level=2)
        temporal = intelligence.get("temporal") or {}
        versions = intelligence.get("version_and_duplicates") or {}
        quality = intelligence.get("ocr_and_parse_quality") or {}
        reuse = intelligence.get("incremental_reuse") or {}
        add_table(
            ["维度", "结果", "边界说明"],
            [
                ["来源时间", temporal.get("source_modified_time_range", "未知"), temporal.get("limitation", "")],
                ["重复/近似版本", "精确组 {}；重复文件 {}；相似簇 {}".format(versions.get("exact_duplicate_groups", 0), versions.get("exact_duplicate_files", 0), versions.get("similar_document_clusters", 0)), versions.get("note", "")],
                ["OCR/解析质量", "Office 图片 OCR {}；截断 {}；失败 {}".format(quality.get("office_embedded_image_ocr_files", 0), quality.get("truncated_text_files", 0), quality.get("failed_files", 0)), "结构化质量分 {}".format(quality.get("structured_average_quality_score", "—"))],
                ["增量复用", "复用检查点 {}；本轮处理 {}".format(reuse.get("reused_parse_checkpoints", 0), reuse.get("newly_processed_files", 0)), "配置与源文件指纹一致时复用"],
            ],
        )
        entities = intelligence.get("entities") or {}
        if entities:
            add_paragraph("实体概览：{}。".format("；".join(
                "{} 已观察 {} 个不同值".format(name, (item or {}).get("observed_distinct_count", 0))
                for name, item in entities.items()
            )))
        add_items(intelligence.get("structured_anomaly_questions"), empty_text="当前结构化数据未形成可靠异常核查问题。")
    doc.add_page_break()
    add_heading("二、全局分类", level=1)
    categories = report.get("global_categories") or []
    classification_coverage = report.get("classification_coverage", {})
    if classification_coverage:
        source_label = {
            "adaptive_analysis_tree": "自适应内容分类树",
            "physical_directory_fallback": "原始目录树（自适应分析不可用时的降级结果）",
            "root_fallback": "扫描根节点（未形成可用分类）",
        }.get(classification_coverage.get("source"), "本地分类结果")
        add_paragraph(
            "分类来源：{}；顶层类别 {} 个；已归入 {} / {} 个已解析文件（{}）。".format(
                source_label,
                classification_coverage.get("top_level_category_count", 0),
                classification_coverage.get("classified_file_count", 0),
                classification_coverage.get("scanned_file_count", classification_coverage.get("parsed_file_count", 0)),
                "完整" if classification_coverage.get("complete") else "需复核未归类或重复归类文件",
            )
        )
    if not categories:
        add_paragraph("暂无可用分类。")
    if categories:
        add_table(
            ["序号", "一级主题", "文件数", "已解析", "状态"],
            [[index, item.get("name", "未命名分类"), item.get("file_count", 0), item.get("parsed_file_count", 0), item.get("classification_status", "classified")]
             for index, item in enumerate(categories[:15], 1)],
        )
    for item in categories[:12]:
        if isinstance(item, dict):
            add_heading(
                "{}（{}：{} 个文件）".format(
                    item.get("name", "未命名分类"), item.get("dimension", "内容类别"), item.get("file_count", 0)
                ),
                level=2,
            )
            add_paragraph(item.get("description") or "暂无分类说明。")
            type_counts = item.get("type_counts", {})
            if type_counts:
                add_paragraph("格式构成：{}。".format(
                    "；".join("{} {}个".format(extension, count) for extension, count in type_counts.items())
                ))
            if item.get("topics"):
                add_paragraph("内容线索：{}。".format("、".join(item["topics"][:12])))
            if item.get("representative_documents"):
                add_paragraph("代表文档：{}".format("；".join(item["representative_documents"][:5])))
            if item.get("conclusion_evidence"):
                add_paragraph("关键结论—证据链：")
                add_conclusion_evidence(item["conclusion_evidence"])
            for subcategory in item.get("subcategories", [])[:5]:
                if not isinstance(subcategory, dict):
                    continue
                add_heading(
                    "{}（{}：{} 个文件）".format(
                        subcategory.get("name", "未命名主题"),
                        subcategory.get("dimension", "内容主题"),
                        subcategory.get("file_count", 0),
                    ),
                    level=3,
                )
                add_paragraph(subcategory.get("description") or "该主题下的文件已完成统一解析。")
                if subcategory.get("topics"):
                    add_paragraph("内容线索：{}。".format("、".join(subcategory["topics"][:10])))
                if subcategory.get("representative_documents"):
                    add_paragraph("代表文档：{}".format("；".join(subcategory["representative_documents"][:3])))
                if subcategory.get("conclusion_evidence"):
                    add_paragraph("关键结论—证据链：")
                    add_conclusion_evidence(subcategory["conclusion_evidence"])
            category_evidence = item.get("evidence_chain", [])
            if category_evidence:
                add_paragraph("分类依据与可回查证据：")
                for evidence in category_evidence[:4]:
                    if not isinstance(evidence, dict):
                        continue
                    location = evidence.get("source_path", "未知文件")
                    if evidence.get("page"):
                        location += "，第 {} 页".format(evidence["page"])
                    if evidence.get("section"):
                        location += "，{}".format(evidence["section"])
                    add_paragraph(
                        "[{}] {}：{}".format(
                            evidence.get("evidence_id") or "元数据证据",
                            location,
                            clean_evidence(evidence.get("text") or evidence.get("fact")),
                        ),
                        style="List Bullet",
                    )
        else:
            add_paragraph(item, style="List Bullet")
    if len(categories) > 12:
        add_paragraph("其余 {} 个低规模类别已保留在系统交互目录和分类清单中，本概览不逐项展开。".format(len(categories) - 12))
    doc.add_page_break()
    add_heading("三、关键发现", level=1)
    add_items(report.get("key_findings"))

    candidates = report.get("direction_candidates") or []
    if candidates:
        add_heading("候选方向排序", level=2)
        add_table(
            ["排名", "候选方向", "综合分", "优先级", "置信度", "证据/独立来源"],
            [[item.get("rank", index), item.get("title", "待命名"), item.get("score", 0), item.get("priority", "—"), item.get("confidence", "—"), "{}/{}".format((item.get("score_breakdown") or {}).get("evidence_count", 0), (item.get("score_breakdown") or {}).get("independent_source_count", 0))]
             for index, item in enumerate(candidates[:5], 1)],
        )
    doc.add_page_break()
    add_heading("四、推荐研究方向", level=1)
    recommendation = report.get("recommended_research_direction") or {}
    add_heading(recommendation.get("title") or "待进一步确定研究方向", level=2)
    lead = add_paragraph()
    lead_run = add_run(lead, "性质：推论；优先级：{}；置信度：{}".format(recommendation.get("priority", "中"), recommendation.get("confidence", "中")))
    set_font(lead_run, size=11, color=(31, 58, 95), bold=True)
    add_paragraph(recommendation.get("rationale") or "需要先生成代表性文档的全文摘要，再确定研究重点。")
    add_heading("拟研究问题", level=3)
    add_items(recommendation.get("research_questions") or recommendation.get("questions"), style="List Number")
    add_heading("建议方法", level=3)
    add_items(recommendation.get("methods"), empty_text="建议采用分层抽样、证据矩阵和专题复核方法。")
    add_heading("方向依据与证据", level=3)
    recommendation_evidence = recommendation.get("evidence_chain", [])
    if recommendation_evidence:
        for evidence in recommendation_evidence[:12]:
            if isinstance(evidence, dict):
                location = evidence.get("source_path", "未知文件")
                if evidence.get("page"):
                    location += "，第 {} 页".format(evidence["page"])
                if evidence.get("section"):
                    location += "，{}".format(evidence["section"])
                add_paragraph("[{}] {}：{}".format(evidence.get("evidence_id") or "元数据证据", location, clean_evidence(evidence.get("text"))), style="List Bullet")
    else:
        add_paragraph("当前未形成可引用正文证据，方向置信度应下调并优先人工复核。")

    doc.add_page_break()
    add_heading("五、其他深入方向建议", level=1)
    for item in report.get("directions", []):
        if isinstance(item, dict):
            add_heading(item.get("direction", "建议"), level=2)
            add_paragraph("类型：{}；置信度：{}".format(item.get("type", "推论"), item.get("confidence", "未标注")))
            for evidence in item.get("evidence_chain", item.get("evidence", [])):
                if isinstance(evidence, dict):
                    add_paragraph("依据：[{}] {}{}".format(
                        evidence.get("evidence_id") or "证据",
                        evidence.get("source_path") or evidence.get("reason", ""),
                        "：" + clean_evidence(evidence.get("text") or evidence.get("fact", "")) if (evidence.get("text") or evidence.get("fact")) else "",
                    ), style="List Bullet")
            if item.get("confidence_note"):
                add_paragraph("置信度说明：{}".format(item["confidence_note"]))
        else:
            add_paragraph(item, style="List Bullet")
    if not report.get("directions"):
        add_paragraph("暂无其他深入方向建议。")
    doc.add_page_break()
    add_heading("六、分析方法与边界", level=1)
    method = report.get("analysis_method", {})
    for key, label in (("parse", "统一解析"), ("deduplication", "精确去重"), ("similarity", "相似聚类"), ("retrieval", "本地证据检索"), ("classification", "自适应分类"), ("traceability", "证据回溯")):
        if method.get(key):
            add_paragraph("{}：{}".format(label, method[key]), style="List Bullet")
    add_paragraph("本报告是数据分析智能体自身的情况概览交付物，不等同于报告整编智能体生成的正式专题报告。模型仅用于语言与研究方向增强；文档解析、OCR、去重、聚类、建树和证据链均在本地执行。")
    doc.save(str(output_path))
